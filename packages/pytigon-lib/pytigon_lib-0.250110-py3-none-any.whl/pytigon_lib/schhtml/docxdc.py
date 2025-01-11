#!/usr/bin/python
# -*- coding: utf-8 -*-
# This program is free software; you can redistribute it and/or modify
# it under the terms of the GNU Lesser General Public License as published by the
# Free Software Foundation; either version 3, or (at your option) any later
# version.
#
# This program is distributed in the hope that it will be useful, but
# WITHOUT ANY WARRANTY; without even the implied warranty of MERCHANTIBILITY
# or FITNESS FOR A PARTICULAR PURPOSE.  See the GNU General Public License
# for more details.

# Pytigon - wxpython and django application framework

# author: "Slawomir Cholaj (slawomir.cholaj@gmail.com)"
# copyright: "Copyright (C) ????/2012 Slawomir Cholaj"
# license: "LGPL 3.0"
# version: "0.1a"

from pytigon_lib.schhtml.basedc import BaseDc, BaseDcInfo
import io

from docx import Document
from docx.shared import Inches
from docx.text.run import Font
from docx.shared import Pt, RGBColor

import docx.enum.text


class DocxDc(BaseDc):
    def __init__(
        self,
        ctx=None,
        calc_only=False,
        width=8.5,
        height=11,
        output_name=None,
        output_stream=None,
        scale=1.0,
        notify_callback=None,
        docx_template_path=None,
        record=False,
    ):
        BaseDc.__init__(
            self,
            calc_only,
            -1,
            -1,
            output_name,
            output_stream,
            scale,
            notify_callback,
            record,
        )
        self.dc_info = DocxDcinfo(self)
        self.type = None

        if width < 0:
            self.width = -1
        if height < 0:
            self.height = 1000000000

        self.last_style_tab = None
        self.handle_html_directly = True
        self.document = Document(docx_template_path)

        self.page_width = width
        self.page_height = height
        self.set_margins(0.5, 0.5, 0.5, 0.5)

        self.map = {
            "body": self.body,
            "p": self.p,
            #            "code": self.p,
            #            "pre": self.p,
            "div": self.div,
            "h1": self.h1,
            "h2": self.h2,
            "h3": self.h3,
            "h4": self.h4,
            "h5": self.h5,
            "h6": self.h6,
            "table": self.table,
            "tr": self.tr,
            "td": self.td,
            "th": self.th,
            # "img": self.image,
        }
        self.last_ = None

        if self.notify_callback:
            self.notify_callback(
                "start",
                {"dc": self},
            )

    def set_margins(self, top, right, bottom, left):
        current_section = self.document.sections[-1]
        if top:
            current_section.top_margin = Inches(top)
        if right:
            current_section.right_margin = Inches(right)
        if bottom:
            current_section.bottom_margin = Inches(bottom)
        if left:
            current_section.left_margin = Inches(left)
        self.body_width = (
            self.page_width
            - current_section.left_margin.inches
            - current_section.right_margin.inches
        )
        self.body_height = (
            self.page_height
            - current_section.top_margin.inches
            - current_section.bottom_margin.inches
        )

    def close(self):
        if self.notify_callback:
            self.notify_callback(
                "end",
                {"dc": self},
            )
        if self.output_stream:
            self.document.save(self.output_stream)
        elif self.output_name:
            self.document.save(self.output_name)

    def annotate(self, what, data):
        if what == "end_tag":
            element = data["element"]
            parent = element.parent
            if element and parent:
                if element.tag in self.map:
                    self.map[element.tag](element, parent)

    def _handle_width_and_height(self, element):
        width = None
        height = None

        if "width" in element.attrs:
            w = element.attrs["width"]
            if "%" in w:
                width = (self.body_width - 0.2) * int(w.replace("%", "")) / 100
            else:
                width = (
                    int(w.replace("px", "").replace("rem", "").replace("em", "")) / 300
                )

        if "height" in element.attrs:
            h = element.attrs["height"]
            if "%" in h:
                height = self.body_height * int(w.replace("%", "")) / 100
            else:
                height = (
                    int(h.replace("px", "").replace("rem", "").replace("em", "")) / 300
                )
        return (width, height)

    # orientation 0 - vertical, 1 - horizontal
    def _inch_from_param(self, param, orientation):
        if "%" in param:
            x = (
                (self.body_height if orientation == 0 else self.body_width)
                * int(param.replace("%", ""))
                / 100
            )
        else:
            x = int(param.replace("px", "").replace("rem", "").replace("em", "")) / 300
        return x

    def _add_style(self, dest_element, source_element):
        if "classes" in source_element.attrs:
            for attr in source_element.attrs["classes"].split(" "):
                if attr.startswith("Style-"):
                    dest_element.style = attr[6:].replace("-", " ")

        padding = None
        margin = None

        if "padding" in source_element.attrs:
            padding = source_element.attrs["padding"].split(" ")
        if "margin" in source_element.attrs:
            margin = source_element.attrs["margin"].split(" ")

        if hasattr(dest_element, "paragraph_format") and (padding or margin):
            dest_element.paragraph_format.left_indent = Inches(0)
            dest_element.paragraph_format.right_indent = Inches(0)
            dest_element.paragraph_format.space_before = Inches(0)
            dest_element.paragraph_format.space_after = Inches(0)

            for tab in (margin, padding):
                if tab:
                    if len(tab) == 4:
                        dest_element.paragraph_format.left_indent += Inches(
                            self._inch_from_param(tab[3], 1)
                        )
                        dest_element.paragraph_format.right_indent += Inches(
                            self._inch_from_param(tab[1], 1)
                        )
                        dest_element.paragraph_format.space_before += Inches(
                            self._inch_from_param(tab[0], 0)
                        )
                        dest_element.paragraph_format.space_after += Inches(
                            self._inch_from_param(tab[2], 0)
                        )
                    elif len(tab) == 2:
                        dest_element.paragraph_format.left_indent += Inches(
                            self._inch_from_param(tab[1], 1)
                        )
                        dest_element.paragraph_format.right_indent += Inches(
                            self._inch_from_param(tab[1], 1)
                        )
                        dest_element.paragraph_format.space_before += Inches(
                            self._inch_from_param(tab[0], 0)
                        )
                        dest_element.paragraph_format.space_after += Inches(
                            self._inch_from_param(tab[0], 0)
                        )
                    else:
                        dest_element.paragraph_format.left_indent += Inches(
                            self._inch_from_param(tab[0], 1)
                        )
                        dest_element.paragraph_format.right_indent += Inches(
                            self._inch_from_param(tab[0], 1)
                        )
                        dest_element.paragraph_format.space_before += Inches(
                            self._inch_from_param(tab[0], 0)
                        )
                        dest_element.paragraph_format.space_after += Inches(
                            self._inch_from_param(tab[0], 0)
                        )
        if "align" in source_element.attrs:
            attr = source_element.attrs["align"]
        else:
            if "text-align" in source_element.attrs:
                attr = source_element.attrs["text-align"]
            else:
                attr = ""
        if attr == "center":
            dest_element.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        else:
            if attr == "right":
                dest_element.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
            else:
                dest_element.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

    def _add_image(self, img, dest_element, width=None, height=None):
        img_stream = io.BytesIO(img)
        dest_element.add_picture(
            img_stream,
            width=Inches(width) if width else None,
            height=Inches(height) if height else None,
        )

    def _any_parent_is_pre(self, element):
        e = element
        while e:
            if hasattr(e, "tag") and e.tag == "pre":
                return True
            e = e.parent
        return False

    def _process_atom_list(self, dest_element, source_element):
        if source_element.atom_list and source_element.atom_list.atom_list:
            for atom in source_element.atom_list.atom_list:
                if atom.style:
                    style = self.dc_info.styles[atom.style]
                    (
                        color,
                        font_family,
                        font_size,
                        font_style,
                        font_weight,
                        text_decoration,
                    ) = style.split(";")
                else:
                    style = None
                if type(atom).__name__ == "BrAtom":
                    dest_element.add_run("\n")
                elif type(atom.data) == str:
                    s = atom.data.replace("»", " ")
                    x = dest_element.add_run(s)
                    if style:
                        if int(font_weight) > 0:
                            x.font.bold = True
                        if int(font_style) == 1:
                            x.font.italic = True
                        x.font.size = Pt(int(int(font_size) * 10 / 100))
                        (r, g, b) = self.rgbfromhex(color)
                        x.font.color.rgb = RGBColor(r, g, b)

                elif type(atom.data).__name__ == "ImgDraw":
                    width, height = self._handle_width_and_height(atom.data.img_tag)
                    self._add_image(
                        atom.data.image,
                        dest_element.add_run(),
                        width,
                        height,
                    )
                else:
                    self._process_atom_list(dest_element, atom.data)

    def h(self, element, level):
        hh = self.document.add_heading("", level)
        self._add_style(hh, element)
        self._process_atom_list(hh, element)

    def p(self, element, parent):
        par = self.document.add_paragraph("")
        self._add_style(par, element)
        self._process_atom_list(par, element)

    def div(self, element, parent):
        par = self.document.add_paragraph("")
        par.paragraph_format.left_indent = 0
        par.paragraph_format.right_indent = 0
        self._add_style(par, element)
        self._process_atom_list(par, element)

    # def image(self, element, parent):
    #    if element.img:
    #        img_stream = io.BytesIO(element.img)
    #        width, height = self._handle_width_and_height(element)
    #        self.document.add_picture(
    #            img_stream,
    #        )

    def body(self, element, parent):
        pass

    def h1(self, element, parent):
        return self.h(element, 0)

    def h2(self, element, parent):
        return self.h(element, 1)

    def h3(self, element, parent):
        return self.h(element, 2)

    def h4(self, element, parent):
        return self.h(element, 3)

    def h5(self, element, parent):
        return self.h(element, 4)

    def h6(self, element, parent):
        return self.h(element, 5)

    def table(self, element, parent):
        tr = element.tr_list
        if len(tr) > 0:
            table = self.document.add_table(rows=len(tr), cols=len(tr[0].td_list))
            self._add_style(table, element)
            i = 0
            for row in tr:
                j = 0
                row_dest = table.rows[i].cells
                for td in row.td_list:
                    try:
                        c = row_dest[j]
                        c._tc.clear_content()
                        p = c.add_paragraph(None)
                        self._add_style(p, td)
                        self._process_atom_list(p, td)
                    except:
                        pass
                    j += 1
                i += 1
        element.tr_list = []

    def tr(self, element, parent):
        if parent.tag == "table":
            parent.tr_list.append(element)
        else:
            parent.parent.tr_list.append(element)

    def td(self, element, parent):
        parent.td_list.append(element)

    def th(self, element, parent):
        parent.td_list.append(element)


class DocxDcinfo(BaseDcInfo):
    def __init__(self, dc):
        BaseDcInfo.__init__(self, dc)

    def get_text_height(self, word, style):
        return 1

    def get_img_size(self, png_data):
        try:
            png_stream = io.BytesIO(png_data)
            image = PIL.Image.open(png_stream)
        except:
            image = None
        if image:
            w, h = image.size
            return (w, h)
        else:
            return (0, 0)
