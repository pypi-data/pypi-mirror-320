from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from functools import partial 
def tableDOC(doc, data, row, col, title):
    row_labels = row
    col_labels = col
    doc.add_paragraph(title)
    table = doc.add_table(rows=1 + len(data), cols=1 + len(col_labels))
    table.style = 'Table Grid'
    header_cell = table.cell(0, 0)
    header_cell.text = "实况\预报"
    header_cell.paragraphs[0].runs[0].font.bold = False
    for idx, label in enumerate(col_labels):
        cell = table.cell(0, idx+1)
        cell.text = label
        cell.paragraphs[0].runs[0].font.bold = False
    for row_idx, (label, row) in enumerate(zip(row_labels, data)):
        table.cell(row_idx + 1, 0).text = label
        table.cell(row_idx + 1, 0).paragraphs[0].runs[0].font.bold = False
        for col_idx, value in enumerate(row):
            table.cell(row_idx + 1, col_idx + 1).text = str(value) 
def add_heading(doc, text):
    heading = doc.add_heading(text, level=1)
    heading.style.font.size = Pt(18)  # 设置字体大小为18磅
def add_text_to_doc(doc, text):
    para = doc.add_paragraph(text)
    para.style = doc.styles['Body Text']
def add_image_to_doc(doc, image_path):
    doc.add_picture(image_path, width=Inches(5.0)) 
def partialFN(tabledata):
    return partial(tableDOC,row=tabledata["labels"], col= tabledata["columns"],title=tabledata["title"])

dataCY ={
    "labels": ["晴", "雨"],
    "columns": ["晴", "雨"],
    "title": f"彩云10min晴雨混淆矩阵"
    } 
dataWTX ={ 
       "labels": ["晴", "雨"],
       "columns": ["晴", "雨"],
       "title": f"维天信10min晴雨混淆矩阵"
       } 
data2CY ={
    "labels": ["晴", "雨"],
    "columns": ["晴", "雨"],
    "title": f"彩云20min晴雨混淆矩阵"
    } 
data2WTX ={ 
       "labels": ["晴", "雨"],
       "columns": ["晴", "雨"],
       "title": f"维天信20min晴雨混淆矩阵"
       } 

# from docx import Document
# doc = Document() 
# doc.save("./test6.docx")