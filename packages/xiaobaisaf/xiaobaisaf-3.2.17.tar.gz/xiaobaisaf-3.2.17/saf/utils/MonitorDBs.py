#! /usr/bin/env python
"""
@Author: xiaobaiTser
@Email : 807447312@qq.com
@Time  : 2023/6/26 21:04
@File  : MonitorDBs.py
"""
import os
try:
    import mysql.connector
except ImportError:
    import os
    os.system('pip install mysql-connector-python')
import sqlite3
# try:
#     import cx_Oracle
# except ImportError:
#     import os
#     os.system('pip install cx_Oracle')
try:
    from docx import Document
except ImportError:
    import os
    os.system('pip install python-docx')


def generate_data_dictionary(database_type, connection_details, save_filename=None):
    """
    依据数据库存储的数据生成数据字典的Word文件
    :param database_type: sqlite、MySQL、Oracle
    :param connection_details:
    # 实例：

    # mysql/oracle样例：
    connection_details = {
        'host': 'localhost',
        'user': 'username',
        'password': 'password',
        'database': 'database_name'
    }
    generate_data_dictionary('mysql', connection_details, 'xiaobai.docx')

    # sqlite样例：
    connection_details = 'D:\\coding\\studentSystem\\xiaobai_db.db'
    generate_data_dictionary('sqlite', connection_details)

    :return:

    """
    # 连接数据库
    if database_type == "mysql":
        connection = mysql.connector.connect(**connection_details)
    elif database_type == "sqlite":
        connection = sqlite3.connect(connection_details)
    # elif database_type == "oracle":
    #     connection = cx_Oracle.connect(connection_details)
    else:
        raise ValueError("Unsupported database type")

    # 进度条
    progress_rate = 100.0

    # 生成数据字典的Word文件
    document = Document()

    # 获取数据库信息
    cursor = connection.cursor()
    if "sqlite" not in database_type:
        if connection_details["database"]:
            databases = [[connection_details["database"]]]
        else:
            cursor.execute("SHOW DATABASES")  # 获取所有数据库
            databases = cursor.fetchall()

        for i, database in enumerate(databases):
            database_name = database[0]
            document.add_heading(database_name, level=1)

            cursor.execute(f"USE {database_name}")  # 切换到当前数据库
            cursor.execute("SHOW TABLES")  # 获取当前数据库的所有数据表
            tables = cursor.fetchall()

            for j, table in enumerate(tables):
                progress_rate = (
                    (i + 1) * (j + 1) / (len(databases) * len(tables))
                ) * 100.0
                table_name = table[0]
                document.add_heading(table_name, level=2)

                # cursor.execute(f"DESCRIBE {table_name}")  # 获取数据表结构
                cursor.execute(f"SHOW FULL COLUMNS from {table_name}")  # 获取数据表结构
                table_structure = cursor.fetchall()

                table_info = document.add_table(rows=1, cols=3)
                table_info.style = "Table Grid"
                table_info.autofit = False

                # 添加表头
                header_cells = table_info.rows[0].cells
                header_cells[0].text = "字段名"
                header_cells[1].text = "数据类型"
                header_cells[2].text = "注释"

                # 添加字段信息
                for field in table_structure:
                    field_name = field[0]
                    data_type = field[1]
                    comment = field[8] if len(field) > 8 else ""
                    row_cells = table_info.add_row().cells
                    row_cells[0].text = field_name
                    row_cells[1].text = data_type
                    row_cells[2].text = comment
                print("", end="\r")
                print("转换进度：%.2f%%" % progress_rate, end="")
        save_filename = f'{connection_details.get("host").replace(".", "_")}_{connection_details.get("database")}'
        # document.save(f'{connection_details.get("host").replace(".", "_")}_{connection_details.get("database")}.docx')
    else:  # sqlite
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")  # 获取所有表名
        tables = cursor.fetchall()

        for i, table in enumerate(tables):
            progress_rate = round(((i + 1) / len(tables)) * progress_rate, 2)
            table_name = table[0]
            document.add_heading(table_name, level=2)

            cursor.execute(f"PRAGMA table_info({table_name})")  # 获取数据表结构
            table_structure = cursor.fetchall()

            table_info = document.add_table(rows=1, cols=3)
            table_info.style = "Table Grid"
            table_info.autofit = False

            # 添加表头
            header_cells = table_info.rows[0].cells
            header_cells[0].text = "字段名"
            header_cells[1].text = "数据类型"
            header_cells[2].text = "注释"

            # 添加字段信息
            for field in table_structure:
                field_name = field[1]
                data_type = field[2]
                comment = field[8] if len(field) > 8 else ""
                row_cells = table_info.add_row().cells
                row_cells[0].text = field_name
                row_cells[1].text = data_type
                row_cells[2].text = comment
            print("", end="\r")
            print("转换进度：%.2f%%" % progress_rate, end="")
        # 保存Word文件
        if not save_filename:
            save_filename = os.path.split(connection_details)[-1].replace(".", "_")
    if save_filename.endswith(".docx"):
        document.save(save_filename)
    else:
        document.save(f"{save_filename}.docx")


# 调用函数
# if __name__ == '__main__':
#     # mysql
#     connection_details = {
#         'host': '192.168.0.240',
#         'port': 3306,
#         'user': 'chandao',
#         'password': '123456',
#         'database': 'chandao'
#     }
#     generate_data_dictionary('mysql', connection_details, 'xiaobai.docx')
